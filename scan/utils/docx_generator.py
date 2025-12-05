from docx import Document
from docx.shared import Pt

def create_word_doc(ai_output):
    doc = Document()
    doc.add_heading("Study Summary", 0)
    doc.add_paragraph(ai_output.summary)
    doc.add_heading("Key Points", level=1)
    for point in ai_output.key_points.split("•")[1:]:
        doc.add_paragraph(point.strip(), style="List Bullet")
    doc.add_heading("Practice Questions", level=1)
    doc.add_paragraph(ai_output.questions)
    return doc